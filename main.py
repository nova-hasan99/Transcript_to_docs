from flask import Flask, request, jsonify, send_file
from docx import Document
import os
import re
import io
import zipfile
import json

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200MB

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|\']+', '', name).strip()[:50]

def format_key(key):
    return key.replace('_', ' ').title()

def format_value(key, value):
    if key == 'videoId':
        return f"https://www.youtube.com/watch?v={value}"
    return value

@app.route('/generate-docs', methods=['POST'])
def generate_docs():
    zip_buffer = io.BytesIO()
    try:
        if 'json_data' not in request.files:
            return jsonify({'error': 'Missing json_data file in form-data'}), 400

        json_file = request.files['json_data']
        json_str = json_file.read().decode('utf-8')

        data = json.loads(json_str)
        if not isinstance(data, list) or len(data) == 0:
            return jsonify({'error': 'Invalid JSON payload. Must be a non-empty list.'}), 400

        raw_channel_name = data[0].get('channelName', 'output_docs')
        zip_name = sanitize_filename(raw_channel_name) or 'output_docs'

        zip_file = zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED)

        used_titles = set()
        count = {}

        for item in data:
            title = item.get('title', 'Untitled')
            safe_title = sanitize_filename(title)

            # ইউনিক নাম গঠন
            if safe_title in used_titles:
                count[safe_title] = count.get(safe_title, 1) + 1
                safe_title = f"{safe_title}_{count[safe_title]}"
            else:
                used_titles.add(safe_title)

            doc = Document()
            doc.add_heading(title, level=1)

            for key, value in item.items():
                doc.add_paragraph(f"{format_key(key)}: {format_value(key, value)}")

            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            zip_file.writestr(f"{safe_title}.docx", doc_bytes.read())
            doc_bytes.close()

        zip_file.close()
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"{zip_name}.zip",
            mimetype='application/zip'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        zip_buffer.close()  # memory leak ঠেকায়

@app.route('/ping', methods=['GET'])
def ping():
    return jsonify({'message': 'API is working!'}), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)
