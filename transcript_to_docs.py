import re
import io
import zipfile
import json
import csv
import requests
from docx import Document
from flask import jsonify, send_file

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|\']+', '', name).strip()[:50]

def format_key(key):
    return key.replace('_', ' ').title()

def format_value(key, value):
    if key == 'videoId' and value:
        return f"https://www.youtube.com/watch?v={value}"
    return value or ''

def download_json_from_url(file_url):
    try:
        response = requests.get(file_url)
        response.raise_for_status()
        return response.content.decode('utf-8')
    except Exception as e:
        raise RuntimeError(f"Failed to download file: {e}")

def generate_zip_from_transcript(request):
    try:
        # Step 1: Check for form field
        file_url = request.form.get('json_url')
        if not file_url:
            return jsonify({'error': 'Missing form-data field: json_url'}), 400

        # Step 2: Download and parse JSON
        try:
            json_str = download_json_from_url(file_url)
            data = json.loads(json_str)
        except Exception as e:
            return jsonify({'error': 'Invalid JSON file from URL', 'details': str(e)}), 400

        if not isinstance(data, list) or len(data) == 0:
            return jsonify({'error': 'Invalid JSON payload. Must be a non-empty list.'}), 400

        # Step 3: Prepare zip name
        raw_channel_name = data[0].get('channelName', 'output_docs')
        zip_name = sanitize_filename(raw_channel_name) or 'output_docs'

        zip_buffer = io.BytesIO()
        zip_file = zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED)

        used_titles = set()
        count = {}

        for item in data:
            title = item.get('title', 'Untitled')
            safe_title = sanitize_filename(title)

            if safe_title in used_titles:
                count[safe_title] = count.get(safe_title, 1) + 1
                safe_title = f"{safe_title}_{count[safe_title]}"
            else:
                used_titles.add(safe_title)

            doc = Document()
            doc.add_heading(title, level=1)

            for key, value in item.items():
                doc.add_paragraph(f"{format_key(key)}: {format_value(key, value)}")

            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            zip_file.writestr(f"{safe_title}.docx", doc_bytes.read())
            doc_bytes.close()

        # ✅ CSV summary
        csv_buffer = io.StringIO()
        fieldnames = data[0].keys()
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(data)
        zip_file.writestr("all_data.csv", csv_buffer.getvalue())
        csv_buffer.close()

        zip_file.close()
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"{zip_name}.zip",
            mimetype='application/zip'
        )

    except Exception as e:
        return jsonify({'error': f"Server Error: {str(e)}"}), 500
